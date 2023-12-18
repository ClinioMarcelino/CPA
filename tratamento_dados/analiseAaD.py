from docx import Document
import pandas as pd
import numpy as np
import difflib
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from docx2pdf import convert


materiasNaoIclusasNaPesquisa = pd.DataFrame(pd.read_excel('C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/MateriasNaoParticipantes.xlsx'))['DISCIPLINA'].tolist()
def verificaDisciplinasValidas(nomeDisciplina):
    if difflib.get_close_matches(nomeDisciplina, materiasNaoIclusasNaPesquisa) == []:
        return False
    return True


def alunosPorTurma():
    df = pd.DataFrame(pd.read_excel('C:/Users/Clinio.freitas/Documents/CPA/AvaliacaoDisciplina2023/ALUNOS.xlsx'))

    codTurmas = df['IDTURMADISC'].tolist()
    cod = df['IDTURMADISC'].unique().tolist()

    lista=[]

    for c in cod:
        codigos = dict()
        codigos['IDTURMADISC'] = c
        codigos['nAlunos'] = codTurmas.count(c)
        lista.append(codigos)

    return pd.DataFrame(lista)

alunoPTurma = alunosPorTurma()
def respostasDisc():
    df = pd.read_excel('C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/qAaDrespostas.xlsx')

    ra = []
    linhas = list()

    for key, values in df.iterrows():
        if values['ra'] not in ra:
            ra.append(values['ra'])
            for i in range(1, 10):
                if not pd.isnull(values[f'cd{i}']):
                    lin = dict()
                    lin[f'codDisc'] = values[f'cd{i}']
                    lin[f'nameDisc'] = values[f'nd{i}']
                    for j in range(0, 19):
                        lin[f'r{j + 1}'] = values[int(f'{(((i - 1) * 19) + j)}')]
                    linhas.append(lin)

    data = pd.DataFrame(linhas)
    data = data.dropna()
    data.replace('Não se Aplica', np.nan, inplace=True)

    disciplinas = dict()
    listaDisc = []

    for key, value in data.iterrows():
        if value['codDisc'] not in listaDisc:
            listaDisc.append(value['codDisc'])
            df = data[data['codDisc'] == value['codDisc']]
            respostas = dict()
            for chave, valor in df.iterrows():
                for k in range(1, 20):
                    # for x in df[f'r{k}'].tolist():
                    a = [x for x in df[f'r{k}'].tolist() if np.isnan(x) == False]
                    respostas[f'r{k}'] = a
            disciplinas[value['codDisc']] = respostas

    a = pd.DataFrame(disciplinas)
    return a

def amplitude(a):
    newlist = [x for x in a if np.isnan(x) == False]
    if len(newlist) > 0:
        return abs(max(newlist)-min(newlist))
    else:
        return 0

def perguntas(i):
    if i == 1:
        return 'Pontualidade e assiduidade'
    elif i == 2:
        return 'Participação nas aulas'
    elif i == 3:
        return 'Dedicação às leituras/atividades e revisão da matéria'
    elif i == 4:
        return 'A apresentação dos critérios de avaliação, da metodologia de ensino e do plano de aula'
    elif i == 5:
        return 'A organização e preparação das aulas'
    elif i == 6:
        return 'A pontualidade e assiduidade'
    elif i == 7:
        return 'A recomendação de bibliografia disponível nas bibliotecas física e virtual'
    elif i == 8:
        return 'A articulação entre teoria e prática nos conteúdos'
    elif i == 9:
        return 'A abordagem de temas atuais e relevantes para a formação profissional'
    elif i == 10:
        return 'O domínio do conteúdo demonstrado pelo professor'
    elif i == 11:
        return 'A clareza dos critérios de avaliação utilizados'
    elif i == 12:
        return 'A didática das aulas'
    elif i == 13:
        return'Uso de metodologias inovadoras'
    elif i == 14:
         return 'A disponibilidade para esclarecer dúvidas'
    elif i == 15:
        return 'A postura ética e profissional no relacionamento com os alunos'
    elif i == 16:
        return 'Leitura e Interpretação de textos'
    elif i == 17:
        return 'Exposição de ideias (escrita e oral)'
    elif i == 18:
        return 'Trabalho em equipe'
    elif i == 19:
        return 'Capacidade de síntese dos conteúdos estudados'

data = pd.read_excel('C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/respostasProfs.xlsx')

ra = []
h = 0
nomes = []
profs = []
linhas = []

for key, values in data.iterrows():
    if values['nome'] not in nomes:
        nomes.append(values['nome'])
        for i in range(1, 8):
            if not pd.isnull(values[f'nd{i}']):
                lin = dict()
                lin['nome'] = values['nome'].title()
                lin['disc'] = values[f'nd{i}']
                for j in range(0, 16):
                    lin[f'rp{j+4}'] = values[int(f'{((i-1)*16)+j}')]
                # print(lin)
                linhas.append(lin)
dataRespProfs = pd.DataFrame(linhas)
dataRespProfs.replace('Não se Aplica', 'ÑA', inplace=True)
# print(dataRespProfs)
# dataRespProfs.to_excel('impossivel.xlsx',index=False)


nomeProfs = []
Professores = []

dfProf = pd.DataFrame(pd.read_excel('C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/PROFESSORES.xlsx'))
disciplinas = respostasDisc()
listaDisc = []

for key,value in dfProf.iterrows():
    if value['PROFESSOR'] not in nomeProfs:
        nomeProfs.append(value['PROFESSOR'])
        for chave, valor in dfProf[dfProf['PROFESSOR'] == value['PROFESSOR']].iterrows():
            if verificaDisciplinasValidas(valor['DISCIPLINA']):
                continue
            Materia = dict()
            Materia['nome'] = value['PROFESSOR'].title()
            Materia['disc'] = valor['DISCIPLINA'].title()
            Materia['IDTURMADIS'] = valor['IDTURMADISC']
            Materia['CURSO'] = valor['CURSO']
            Materia['CODDISC'] = valor['CODDISC']
            Materia['TURNO'] = valor['TURNO'].title()
            Materia['TURMA'] = valor['TURMA']
            for k,v in dataRespProfs[dataRespProfs['nome']==Materia['nome']].iterrows():
                if v['disc'].title() == valor['DISCIPLINA'].title():
                    for i in range(4,20):
                        Materia[f'rp{i}'] = v[f'rp{i}']
            Professores.append(Materia)

Prof = pd.DataFrame(Professores)
# Prof = Prof.dropna()
Prof.replace(np.nan,'S/R')
print(Prof)

professores = Prof['nome'].unique().tolist()

for x in professores:
    doc = Document()

    arquivoName = f'C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/devolutivasProfs/{x}.docx'

    doc.add_picture(
        'C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/capa-relatorio-avaliacao-disciplinas.png',
        width=Inches(6))

    doc.add_picture('C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/excel/CPAlogo.png', width=Inches(1.5))
    p1 = doc.add_paragraph(f'Prezado(a) Professor(a)')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph(f'A Comissão Própria de Avaliação da UniLaSalle-RJ apresenta o relatório individualizado com os resultados da Pesquisa de Avaliação de Disciplinas (PAD) referente ao segundo semestre de 2023. ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('A pesquisa foi realizada entre 08 e 29 de novembro. Nesta edição, contamos com a participação de 35% dos estudantes e 83% dos docentes da instituição.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('Devido ao contínuo desenvolvimento institucional, a CPA reformulou sua política de autoavaliação em 2022. Essa nova orientação está em vigor durante o período do Plano de Desenvolvimento Institucional (PDI) 2022-2026 e baseia-se em cinco objetivos norteadores, além de um conjunto de novos instrumentos de avaliação que possuem critérios metodológicos explícitos e garantias aprimoradas de proteção de dados.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('A PAD é um desses instrumentos e foi aplicada pela primeira vez em maio de 2023. Nesta segunda edição, o questionário foi aprimorado atendendo a reflexões e comentários dos professores e coordenadores debatidas pela equipe da CPA.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('É importante ressaltar que um dos principais objetivos da nova política de autoavaliação é disseminar uma cultura de avaliação participativa, propositiva e não punitiva. Por meio de estratégias diferenciadas e incrementais, procuramos aumentar a participação da comunidade acadêmica no processo avaliativo, tanto alcançando um maior número de respondentes nas pesquisas quanto promovendo o conhecimento e a análise da instituição como um todo, por meio da constante divulgação dos resultados das pesquisas. Isso culminará em propostas de medidas para o aperfeiçoamento institucional.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('Nesse sentido, a CPA está divulgando individualmente para cada professor a avaliação de suas turmas, fornecendo todas as informações necessárias: tanto a média dos respondentes, quanto duas medidas de variabilidade. Dessa forma, busca-se o conhecimento sobre a avaliação das disciplinas lecionadas e o aprimoramento profissional.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('Neste relatório, você receberá também a média de autoavaliação dos alunos. Para cada disciplina, o discente avaliou sua pontualidade e assiduidade, sua participação nas aulas e dedicação aos estudos. ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('A CPA solicita especial atenção ao número ainda baixo de discentes respondentes nesta edição da pesquisa. Em alguns casos, não é possível considerar uma análise estatística dos resultados. No entanto, a CPA considera fundamental o envio dos resultados de todos os professores, a fim de permitir-lhes conhecer os resultados mesmo das disciplinas em que os alunos não participaram em peso da avaliação.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('Esperamos contar com o envolvimento de todos vocês na divulgação e sensibilização dos alunos para a próxima edição da pesquisa em 2024.1. Aumentar o número de respondentes é fundamental para consolidar a legitimidade e a validade deste levantamento e para a disseminação da cultura de participação e comprometimento com a avaliação institucional como um todo.')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1 = doc.add_paragraph('Valorizar a autoavaliação contínua para o aprimoramento institucional é responsabilidade de todos nós!')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    atenciosamente = doc.add_paragraph(f'Atenciosamente,',)
    atenciosamente.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    assinatura = doc.add_paragraph(f'Denise Salles\nPresidente da CPA')
    assinatura.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    for key,value in Prof[Prof['nome'] == x].iterrows():
        if verificaDisciplinasValidas(value['disc']):
            continue
        uau = alunoPTurma[alunoPTurma['IDTURMADISC']==value["IDTURMADIS"]]
        doc.add_paragraph(f'Disciplina : {value["disc"]}', style='List Bullet')
        doc.add_paragraph(f'Cod.Disciplina: {value["IDTURMADIS"]}', style='List Bullet')
        doc.add_paragraph(f'Turno: {value["TURNO"]}', style='List Bullet')
        doc.add_paragraph(f'Turma: {value["TURMA"]}', style='List Bullet')
        doc.add_paragraph(f'Nº de Alunos: {max(uau["nAlunos"].values)}', style='List Bullet')
        tam = 0
        try:
            tam = len(disciplinas[value["IDTURMADIS"]][f"r1"])
            doc.add_paragraph(f'Total de respostas: {tam}', style='List Bullet')
            doc.add_paragraph(f'% de respondentes: {(tam / max(uau["nAlunos"].values)) * 100:.2f}%',
                              style='List Bullet')
            doc.add_paragraph('Itens de Auto Avaliação dos Alunos (Média da Turma):')
            doc.add_paragraph(f'Pontualidade e assiduidade: {np.nanmean(disciplinas[value["IDTURMADIS"]][f"r1"]):.2f}',
                              style='List Bullet')
            doc.add_paragraph(f'Participação nas aulas: {np.nanmean(disciplinas[value["IDTURMADIS"]][f"r2"]):.2f}',
                              style='List Bullet')
            doc.add_paragraph(
                        f'Dedicação às leituras/atividades e revisão da matéria: {np.nanmean(disciplinas[value["IDTURMADIS"]][f"r3"]):.2f}',
                        style='List Bullet')

            for i in range(4, 20):
                if i == 4:
                    table = doc.add_table(rows=1, cols=6, style="Table Grid")
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[0].width = Cm(20)
                    hdr_cells = table.rows[0].cells

                    hdr_cells[0].text = 'Itens Avaliação da Disciplina'
                    hdr_cells[1].text = 'Auto Avaliação Professor'
                    hdr_cells[2].text = 'Média Alunos'
                    hdr_cells[3].text = 'Desvio Padrão Alunos'
                    hdr_cells[4].text = 'Amplitude Alunos'
                    hdr_cells[5].text = '% N/A'

                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = f'{np.nanmean(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[3].text = f'{np.nanstd(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[4].text = f'{amplitude(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[5].text = '-'

                elif i < 16:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = f'{np.nanmean(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[3].text = f'{np.nanstd(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[4].text = f'{amplitude(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[5].text = '-'
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = f'{np.nanmean(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[3].text = f'{np.nanstd(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[4].text = f'{amplitude(disciplinas[value["IDTURMADIS"]][f"r{i}"]):.2f}'
                    row_cells[
                        5].text = f'{(abs(len(disciplinas[value["IDTURMADIS"]]["r1"]) - len(disciplinas[value["IDTURMADIS"]][f"r{i}"])) / tam) * 100:.1f}%'


        except:
            tam = 0
            doc.add_paragraph(f'Total de respostas: {tam}', style='List Bullet')
            doc.add_paragraph(f'% de respondentes: {(tam / max(uau["nAlunos"].values)) * 100:.2f}%',
                              style='List Bullet')
            doc.add_paragraph('Itens de Auto Avaliação dos Alunos (Média da Turma):')
            doc.add_paragraph(f'Pontualidade e assiduidade: Sem respondentes',
                              style='List Bullet')
            doc.add_paragraph(f'Participação nas aulas: Sem respondentes',
                              style='List Bullet')
            doc.add_paragraph(
                f'Dedicação às leituras/atividades e revisão da matéria: Sem respondentes',
                style='List Bullet')

            for i in range(4, 20):
                if i == 4:
                    table = doc.add_table(rows=1, cols=6, style="Table Grid")
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[0].width = Cm(20)
                    hdr_cells = table.rows[0].cells

                    hdr_cells[0].text = 'Itens Avaliação da Disciplina'
                    hdr_cells[1].text = 'Auto Avaliação Professor'
                    hdr_cells[2].text = 'Média Alunos'
                    hdr_cells[3].text = 'Desvio Padrão Alunos'
                    hdr_cells[4].text = 'Amplitude Alunos'
                    hdr_cells[5].text = '% N/A'

                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = '-'
                    row_cells[3].text = '-'
                    row_cells[4].text = '-'
                    row_cells[5].text = '-'

                elif i < 16:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = '-'
                    row_cells[3].text = '-'
                    row_cells[4].text = '-'
                    row_cells[5].text = '-'
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'{perguntas(i)}:'
                    row_cells[1].text = f"{value[f'rp{i}']}"
                    row_cells[2].text = f'-'
                    row_cells[3].text = f'-'
                    row_cells[4].text = f'-'
                    row_cells[5].text = f'-'
            print(f'Turma sem respondentes {value["disc"]}')

        doc.add_page_break()

    # doc.add_paragraph('A CPA, junto com todos, faz uma UNILASALLE cada vez melhor!')
    doc.save(arquivoName)

    saida = f'C:/Users/Clinio.freitas/PycharmProjects/pythonCPA2023-2/devolutivasProfs/emPDF/{x}.pdf'

    # CONVERSÃO PARA PDF
    convert(arquivoName, saida)

    print(f'Arqvivo prof {x} salvo')

